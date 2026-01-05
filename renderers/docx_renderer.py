from docx import Document
from docx.shared import Pt, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ROW_HEIGHT_RULE

def hex_to_rgb(hex_str):
    if not hex_str: return RGBColor(0, 0, 0)
    hex_str = hex_str.lstrip('#')
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def get_theme_color(theme, color_key):
    """Resolves 'primary_color' to actual hex, or checks if it's already hex."""
    if color_key in theme:
        return theme[color_key]
    return color_key # logic to handle direct hex codes if passed

class DocxRenderer:
    """
    Renders content into a Microsoft Word (.docx) document using python-docx.
    Manually handles layout, fonts, and table structures to mimic the design.
    """
    def get_font_size(self, size_key, default_pt=10):
        """
        Resolves font size from theme['typography']['docx'] -> theme['typography']['default']
        Returns a Pt object.
        size_key example: 'base', 'h1', 'h2', 'small', 'footer'
        """
        typography = self.theme.get('typography', {})
        docx_typo = typography.get('docx', {})
        default_typo = typography.get('default', {})
        
        full_key = f"font_size_{size_key}"
        
        # 1. Check DOCX specific
        if full_key in docx_typo:
            return Pt(docx_typo[full_key])
            
        # 2. Check Default
        if full_key in default_typo:
            return Pt(default_typo[full_key])
            
        # 3. Fallback
        return Pt(default_pt)

    def __init__(self, theme):
        self.theme = theme
        self.doc = Document()
        self.setup_page_layout()

    def setup_page_layout(self):
        """Sets up global page margins based on style.yaml configuration."""
        sections = self.doc.sections
        margins = self.theme.get('margins', {})
        for section in sections:
            section.top_margin = Mm(margins.get('top', 12.7))
            section.bottom_margin = Mm(margins.get('bottom', 12.7))
            section.left_margin = Mm(margins.get('left', 12.7))
            section.right_margin = Mm(margins.get('right', 12.7))

    def save(self, output_path):
        self.doc.save(output_path)
        print(f"Saved DOCX to: {output_path}")

    def render(self, content_data):
        """
        Main Render Loop:
        Iterates through sections and calls specific render methods based on 'type'.
        Handles manual 'page_break_before' logic by inserting Section Breaks.
        """
        # 1. Sections
        for section in content_data.get('sections', []):
            block_type = section.get('type')
            config = section.get('config', {})
            
            # Handle Page Break
            if config.get('page_break_before'):
                # Use Section Break to reset margins (prevent 0mm top margin bleed)
                self.doc.add_section(WD_SECTION.NEW_PAGE)
                new_section = self.doc.sections[-1]
                
                # Reset margins to theme defaults
                margins = self.theme.get('margins', {})
                new_section.top_margin = Mm(margins.get('top', 12.7))
                new_section.bottom_margin = Mm(margins.get('bottom', 12.7))
                new_section.left_margin = Mm(margins.get('left', 12.7))
                new_section.right_margin = Mm(margins.get('right', 12.7))
                
                # Since we added a section, we don't need add_page_break
            
            if block_type == 'stripe_block':
                 self.render_stripe(config)
            elif block_type == 'header_block':
                self.render_header_block(config)
            elif block_type == 'section_title_block':
                self.render_section_title_block(config)
            elif block_type == 'compound_text_block':
                self.render_compound_text_block(config)
            elif block_type == 'text_block':
                self.render_text_block(config)
            elif block_type == 'grid_block':
                self.render_grid_block(config)
            elif block_type == 'list_block':
                self.render_list_block(config)
            elif block_type == 'plain_list_block':
                self.render_plain_list_block(config)
            elif block_type == 'compact_list_block':
                self.render_compact_list_block(config)
            elif block_type == 'text_grid_block':
                self.render_text_grid_block(config)
            elif block_type == 'project_block':
                self.render_project_block(config)
            else:
                print(f"Warning: Unknown block type '{block_type}'")
            
        # 2. Apply Footer (Global)
        footer_config = content_data.get('config', {}).get('footer')
        if footer_config:
            self.apply_footer(footer_config)

    def create_element(self, name):
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        element.set(qn(name), value)

    def add_page_number(self, paragraph):
        # Page
        run = paragraph.add_run()
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
        
        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def add_num_pages(self, paragraph):
        # NumPages
        run = paragraph.add_run()
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
        
        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "NUMPAGES"
        
        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def apply_footer(self, footer_config):
        text = footer_config.get('text', '')
        show_pages = footer_config.get('show_pages', False)
        theme_footer = self.theme.get('footer', {})
        font_size = self.get_font_size('footer', 8)
        color_hex = theme_footer.get('text_color', '#666666').replace('#', '')
        
        def _set_footer_content(footer_obj):
            footer_obj.is_linked_to_previous = False 
            
            # Clear existing
            for p in footer_obj.paragraphs:
                p._element.getparent().remove(p._element)
                
            p = footer_obj.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Footer Text
            if text:
                run = p.add_run(text)
                run.font.size = font_size
                run.font.color.rgb = RGBColor.from_string(color_hex)
                
            # Page Numbers
            if show_pages:
                if text:
                    run = p.add_run(" | ")
                    run.font.size = font_size
                    run.font.color.rgb = RGBColor.from_string(color_hex)
                
                run = p.add_run("Page ")
                run.font.size = font_size
                run.font.color.rgb = RGBColor.from_string(color_hex)
                self.add_page_number(p)
                
                run = p.add_run(" of ")
                run.font.size = font_size
                run.font.color.rgb = RGBColor.from_string(color_hex)
                self.add_num_pages(p)

        # Iterate all sections
        for section in self.doc.sections:
            _set_footer_content(section.footer)
            
            if section.different_first_page_header_footer:
                _set_footer_content(section.first_page_footer)

    # --- BLOCK RENDERERS ---

    def render_stripe(self, config):
        # STRATEGY: Floating Table (Absolute Position)
        # 1. Restore Standard Margins (so Page 2+ works normally)
        section = self.doc.sections[0]
        margins = self.theme.get('margins', {})
        section.top_margin = Mm(margins.get('top', 12.7))
        section.header_distance = Mm(12.7) # Standard
        
        # 2. Keep Headers distinct (Page 1 Empty, Page 2 Standard)
        section.different_first_page_header_footer = True
        
        # Clear First Page Header (Ensure empty)
        fp_header = section.first_page_header
        for p in fp_header.paragraphs:
            p._element.getparent().remove(p._element)
            
        # Clear Primary Header (Page 2+) of any previous spacer hacks
        header = section.header
        for p in header.paragraphs:
            # We only remove if it looks like our spacer (font size 36)
            # Actually, safe to just clear for now as we don't have other header content defined yet.
            if len(p.runs) > 0 and p.runs[0].font.size == Pt(36):
                 p._element.getparent().remove(p._element)

        # 3. Render Stripe in BODY with FLOATING POSITION
        # anchored to "page" (physical edge), y=0.
        
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False
        
        # Table Position Properties (tblpPr)
        tblPr = table._tbl.tblPr
        tblpPr = OxmlElement('w:tblpPr')
        tblpPr.set(qn('w:vertAnchor'), 'page')
        tblpPr.set(qn('w:horzAnchor'), 'page')
        tblpPr.set(qn('w:tblpY'), '0') # Top edge
        tblpPr.set(qn('w:tblpX'), '0') # Left edge
        # Wrapping 'around' ensures it floats
        tblpPr.set(qn('w:tblpYSpec'), 'top') # Force Top
        tblpPr.set(qn('w:tblpXSpec'), 'center') # Center Horizontally
        
        tblPr.append(tblpPr)
        
        # Width: Full Page Width
        # Letter is 215.9mm. Let's make it slightly wider to be safe (bleed).
        page_width = Mm(220).twips
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), str(page_width))
        tblW.set(qn('w:type'), 'dxa')
        
        # Fill Color
        cell = table.cell(0,0)
        c_val = config.get('color')
        hex_color = get_theme_color(self.theme, c_val).lstrip('#')
        
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)
        
        # Exact Height (Thin)
        row = table.rows[0]
        row.height = Pt(8) 
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        # Content (Empty/Tiny)
        p_cell = cell.paragraphs[0]
        p_cell.paragraph_format.space_before = Pt(0)
        p_cell.paragraph_format.space_after = Pt(0)
        run_cell = p_cell.add_run()
        run_cell.font.size = Pt(1) 
        
        # Remove the extra paragraph added by Word after a table?
        # Actually, since it's floating, the text flowing "around" might start nicely.
        # But we added it to the body stream.
        # Let's add a small spacer paragraph below to ensure name doesn't overlap 
        # if float interaction is weird (though y=0 should be far above margin).
        # Standard margin is 12.7mm. Stripe is 3mm. No overlap.
        pass
    def render_header_block(self, config):
        # Title (Name)
        title = config.get('title', '')
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.name = self.theme.get('font_header', 'Arial')
        run.font.size = Pt(28) # Configurable?
        run.font.color.rgb = hex_to_rgb(get_theme_color(self.theme, 'primary_color'))

        # Subtitle
        subtitle = config.get('subtitle')
        if subtitle:
            p_sub = self.doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sub = p_sub.add_run(subtitle)
            run_sub.font.size = Pt(11)
            run_sub.font.color.rgb = RGBColor(100, 100, 100) # Grey

    def render_section_title_block(self, config):
        content = config.get('content', '')
        style = config.get('style', 'normal')
        self.add_section_title(content, style=style)

    def render_compound_text_block(self, config):
        p = self.doc.add_paragraph()
        
        # Alignment
        alignment = config.get('font_alignment', 'center').lower()
        if alignment == 'left': p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        separator = config.get('separator', ' • ')
        default_color = config.get('font_color', 'text_color')
        default_size = config.get('font_size', 10)
        
        items = config.get('items', [])
        for i, item in enumerate(items):
            # Add separator if not first
            if i > 0:
                r_sep = p.add_run(separator)
                r_sep.font.size = Pt(default_size)
                # Resolve separator color? Use default
            
            text = item.get('text', '')
            link = item.get('link')
            
            # Item specific style or default
            color_key = item.get('font_color', default_color)
            color_hex = get_theme_color(self.theme, color_key)
            
            if link:
                self.add_hyperlink(p, link, text, color=color_hex)
            else:
                r = p.add_run(text)
                r.font.size = Pt(default_size)
                r.font.color.rgb = hex_to_rgb(color_hex)

    def render_text_block(self, config):
        content = config.get('content', '')
        style = config.get('style', 'normal')
        
        if style == 'shaded':
            # Shaded Box Logic
            self.doc.add_paragraph().paragraph_format.space_after = Pt(6) # Buffer
            table = self.doc.add_table(rows=1, cols=1) # Width removed
            # Remove table defaults?
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Mm(190) # Full width
            
            cell = table.cell(0,0)
            
            # Shading
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), "F2F2F2") 
            tcPr.append(shd)
            
            # PADDING (Margins)
            # 15px ~= 11pt ~= 220 dxa (1/20 pt)
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom', 'left', 'right']:
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), '220')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

            # Border (Left Accent) - Always apply for shaded style
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '64') # 8pt (thick)
            left.set(qn('w:space'), '0')
            
            # Resolve accent color
            c_hex = get_theme_color(self.theme, config.get('border_color', 'accent_color'))
            left.set(qn('w:color'), c_hex.lstrip('#'))
            tcBorders.append(left)

            p = cell.paragraphs[0]
            # Remove spacing inside the box, let padding handle it
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.2 # slightly tighter inside box
            
            self.add_markdown_text(p, content)

            # Add spacing AFTER the table
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(12)
            
        elif style == 'left_border':
            # Left Border Logic (No Shading)
            self.doc.add_paragraph().paragraph_format.space_after = Pt(6) 
            table = self.doc.add_table(rows=1, cols=1) 
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Mm(190) 
            
            cell = table.cell(0,0)
            tcPr = cell._tc.get_or_add_tcPr()
            
            # PADDING (Left only mostly, but let's give top/bottom slight)
            tcMar = OxmlElement('w:tcMar')
            for side, width in [('top', '0'), ('bottom', '0'), ('left', '220'), ('right', '0')]:
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), width)
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

            # Border (Left Primary)
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '32') # 4pt (Standard for this look)
            left.set(qn('w:space'), '0')
            c_hex = get_theme_color(self.theme, 'primary_color')
            left.set(qn('w:color'), c_hex.lstrip('#'))
            tcBorders.append(left)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            # Support HTML color spans in markdown (basic parser update needed or handle here?)
            # Our add_markdown_text doesn't handle HTML strings.
            # We will rely on user providing separate logic OR we upgrade add_markdown_text to handle simple <span>?
            # For now, let's proceed and if color fails we fix.
            self.add_markdown_text(p, content)

            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(12)
            
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(12)
            
        elif style == 'shaded_primary':
             # Left Border + Shading (Blue Border)
            self.doc.add_paragraph().paragraph_format.space_after = Pt(6) 
            table = self.doc.add_table(rows=1, cols=1) 
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Mm(190) 
            
            cell = table.cell(0,0)
            tcPr = cell._tc.get_or_add_tcPr()
            
            # Padding
            tcMar = OxmlElement('w:tcMar')
            for side, width in [('top', '80'), ('bottom', '80'), ('left', '220'), ('right', '80')]:
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), width)
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

            # Shading
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F2F2F2')
            tcPr.append(shd)

            # Border (Left Primary)
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '32')
            left.set(qn('w:space'), '0')
            c_hex = get_theme_color(self.theme, 'primary_color')
            left.set(qn('w:color'), c_hex.lstrip('#'))
            tcBorders.append(left)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            self.add_markdown_text(p, content)

            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(12)
            
        else:
            p = self.doc.add_paragraph()
            self.add_markdown_text(p, content)

    def render_grid_block(self, config):
        # 2-Column Grid (using Table)
        items = config.get('items', [])
        default_cols = len(items) if items else 2
        columns = config.get('columns', default_cols)
        
        # Check for Shaded Style
        is_shaded = config.get('style') == 'shaded'
        if is_shaded:
            # Create Container Table for Border
            container = self.doc.add_table(rows=1, cols=1)
            container.autofit = False
            container.allow_autofit = False
            container.columns[0].width = Mm(190) # Full width approx
            
            cell_container = container.cell(0,0)
            target = cell_container
            
            # Apply Left Border & Padding to Container
            tcPr = cell_container._tc.get_or_add_tcPr()
            
            # Padding
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom', 'left', 'right']:
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), '220') # 15px
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
            
            # Border
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '64') # 8pt
            left.set(qn('w:space'), '0')
            c_hex = get_theme_color(self.theme, 'accent_color')
            left.set(qn('w:color'), c_hex.lstrip('#'))
            tcBorders.append(left)
            
        else:
            target = self.doc

        # Section Title if present
        if config.get('title'):
            self.add_section_title(config.get('title'), target=target)

        # Grid Table
        table = target.add_table(rows=1, cols=columns)
        table.autofit = False
        table.allow_autofit = False
        
        for i, col_data in enumerate(items):
            # Calculate row and column index
            row_idx = i // columns
            col_idx = i % columns
            
            # Add new row if needed
            if row_idx >= len(table.rows):
                table.add_row()
                
            cell = table.cell(row_idx, col_idx)
            
            # Header
            header = col_data.get('header')
            if header:
                p = cell.paragraphs[0]
                run = p.add_run(header)
                run.bold = True
                run.font.color.rgb = hex_to_rgb(get_theme_color(self.theme, 'primary_color'))
                run.font.size = Pt(11)
            else:
                p = cell.paragraphs[0] # Empty start
                
            # Content List
            content_list = col_data.get('content', [])
            for line in content_list:
                p_item = cell.add_paragraph()
                p_item.style = 'List Bullet' # Use built-in bullet style
                self.add_markdown_text(p_item, line)

    def render_list_block(self, config):
        title = config.get('title')
        if title:
            self.add_section_title(title)
            
        style = config.get('style', 'simple')
        items = config.get('items', [])
        
        # We treat all styles similarly if they have structured data
        for item in items:
            # Check if we have left/right text structure
            if item.get('left_text') or item.get('right_text'):
                # 1. Headline Row (Table for Left/Right alignment)
                table = self.doc.add_table(rows=1, cols=2)
                table.autofit = False
                
                # Widths (75% / 25% approx)
                # Total Page Body: 210mm (A4/Letter width approx) - 25.4mm (12.7mm * 2 margins) = ~184.6mm
                # We target 184mm total table width to fill the printable area.
                # Left Column (Title): 138mm (~75%)
                # Right Column (Date): 46mm (~25%)
                # This ensures long job titles don't wrap prematurely while keeping dates aligned.
                w_left = Mm(138)
                w_right = Mm(46)
                
                table.columns[0].width = w_left
                table.columns[1].width = w_right
                
                # Force cell widths for first row explicitly (helps Word rendering engine respect widths)
                table.cell(0, 0).width = w_left
                table.cell(0, 1).width = w_right
                
                # Left: Job Title
                c1 = table.cell(0, 0)
                p1 = c1.paragraphs[0]
                r1 = p1.add_run(item.get('left_text', ''))
                r1.bold = True
                r1.font.size = Pt(11)
                
                # Right: Dates
                c2 = table.cell(0, 1)
                p2 = c2.paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r2 = p2.add_run(item.get('right_text', ''))
                r2.bold = True
                r2.font.size = Pt(10) # Reduced to fit in narrower column
                r2.font.color.rgb = hex_to_rgb(get_theme_color(self.theme, 'accent_color'))
                
                # Subtext (Company)
                if item.get('sub_text'):
                    p_sub = self.doc.add_paragraph()
                    p_sub.paragraph_format.space_after = Pt(2)
                    run_sub = p_sub.add_run(item.get('sub_text'))
                    run_sub.italic = True
                    run_sub.font.color.rgb = hex_to_rgb("#666666") 
            else:
                # Fallback for simple string items
                pass
            
            # Details (Bullets)
            for detail in item.get('details', []):
                p_det = self.doc.add_paragraph()
                p_det.style = 'List Bullet'
                self.add_markdown_text(p_det, detail)
                
            # Spacing between jobs
            self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def render_plain_list_block(self, config):
        title = config.get('title')
        if title:
            # Check for accented style
            style = config.get('title_style', 'normal')
            self.add_section_title(title, style=style)

        items = config.get('items', [])
        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            
            text_content = item
            is_small = False
            
            if isinstance(item, dict):
                text_content = item.get('text', '')
                if item.get('style') == 'small':
                    is_small = True
            
            # Helper to add runs with potential override
            # We need to modify add_markdown_text to accept font overrides or handle it here
            # Since add_markdown_text is simple, let's just make a new one or inline logic
            # for bolding etc. BUT we want to apply small size to everything.
            
            # Let's temporarily override paragraph style or use a custom add_markdown
            self.add_markdown_text(p, text_content, font_size_override=Pt(7) if is_small else None, color_override=RGBColor(100,100,100) if is_small else None)

    def render_compact_list_block(self, config):
        title = config.get('title')
        if title:
            style = config.get('title_style', 'normal')
            self.add_section_title(title, style=style)

        items = config.get('items', [])
        for item in items:
            # Table for layout: Left (Content) | Right (Date)
            table = self.doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # Widths (Approx 80/20)
            # Assuming ~190mm page width
            table.columns[0].width = Mm(150)
            table.columns[1].width = Mm(40)
            
            c1 = table.cell(0, 0)
            c2 = table.cell(0, 1)
            
            # Content
            p1 = c1.paragraphs[0]
            self.add_markdown_text(p1, item.get('content', ''))
            
            # Date
            p2 = c2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p2.add_run(item.get('date', ''))
            
            # Add dotted bottom border to cells
            for cell in [c1, c2]:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'dotted')
                bottom.set(qn('w:sz'), '4') # 1/2 pt
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), 'CCCCCC')
                tcBorders.append(bottom)
            
            # Small spacing after table?
            # docx tables don't have margin-bottom. 
            # We add a small empty paragraph
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.line_spacing = Pt(2)
            spacer.paragraph_format.space_after = Pt(2)

    def render_text_grid_block(self, config):
        title = config.get('title')
        if title:
            style = config.get('title_style', 'normal')
            self.add_section_title(title, style=style)
            
        columns = config.get('columns', 2)
        items = config.get('items', [])
        style = config.get('style', 'simple')
        
        # Container Logic (Shading/Border)
        target = self.doc
        if style in ['shaded', 'shaded_primary', 'left_border']:
             # Create Container 1x1
             container = self.doc.add_table(rows=1, cols=1)
             container.autofit = False
             container.allow_autofit = False
             container.columns[0].width = Mm(190)
             
             cell_container = container.cell(0,0)
             tcPr = cell_container._tc.get_or_add_tcPr()
             
             # Resolve Border Color
             # Default depends on style
             if style == 'shaded':
                 def_color = 'accent_color'
                 fill = 'F2F2F2'
             elif style == 'shaded_primary':
                 def_color = 'primary_color'
                 fill = 'F2F2F2'
             else: # left_border
                 def_color = 'primary_color'
                 fill = 'auto'

             # Override?
             if config.get('border_color'):
                 color_key = config.get('border_color')
                 resolved = get_theme_color(self.theme, color_key)
                 if resolved: c_hex = resolved.lstrip('#')
                 else: c_hex = color_key.lstrip('#')
             else:
                 c_hex = get_theme_color(self.theme, def_color).lstrip('#')

             # Padding (Tight for all these styles per user preference)
             if style == 'left_border':
                 # No top/bottom padding margin for left_border usually, or minimal
                 pads = [('top','0'), ('bottom','0'), ('left','220'), ('right','0')]
             else:
                 # Shaded box - tight padding as requested "gain size"
                 pads = [('top','40'), ('bottom','40'), ('left','220'), ('right','80')]
                 
             tcMar = OxmlElement('w:tcMar')
             for side, width in pads:
                 node = OxmlElement(f'w:{side}')
                 node.set(qn('w:w'), width)
                 node.set(qn('w:type'), 'dxa')
                 tcMar.append(node)
             tcPr.append(tcMar)
             
             # Fill
             if fill != 'auto':
                 shd = OxmlElement('w:shd')
                 shd.set(qn('w:val'), 'clear')
                 shd.set(qn('w:color'), 'auto')
                 shd.set(qn('w:fill'), fill)
                 tcPr.append(shd)
                 
             # Border (Left)
             tcBorders = tcPr.find(qn('w:tcBorders'))
             if tcBorders is None:
                 tcBorders = OxmlElement('w:tcBorders')
                 tcPr.append(tcBorders)
                 
             left = OxmlElement('w:left')
             left.set(qn('w:val'), 'single')
             left.set(qn('w:sz'), '32') # 4pt
             left.set(qn('w:space'), '0')
             left.set(qn('w:color'), c_hex)
             tcBorders.append(left)
             
             # Remove default paragraph in container
             cell_container.paragraphs[0]._element.getparent().remove(cell_container.paragraphs[0]._element)
             
             # Target is now the cell to add the grid table into
             # Nesting tables for layout
             target = cell_container

        # Inner Grid Table
        table = target.add_table(rows=1, cols=columns)
        table.autofit = False
        table.allow_autofit = False # Important for nested alignment
        
        # Approximate Widths (If nested, 100% of cell)
        # If nested, available width is ~180mm (190 - padding)
        # 2 cols = ~90mm
        col_width = Mm(90 if style != 'simple' else 95)
        for i in range(columns):
            table.columns[i].width = col_width

        for i, col_data in enumerate(items):
            # Calculate row and column index
            row_idx = i // columns
            col_idx = i % columns
            
            # Add new row if needed
            if row_idx >= len(table.rows):
                table.add_row()
                
            cell = table.cell(row_idx, col_idx)
            
            # Content List (paragraphs, no bullets)
            content_list = col_data.get('content', [])
            
            # Clear default empty paragraph if we are adding our own, 
            # or use it for the first item.
            p = cell.paragraphs[0]
            
            for idx, line in enumerate(content_list):
                if idx == 0:
                    current_p = p
                else:
                    current_p = cell.add_paragraph()
                
                current_p.paragraph_format.space_after = Pt(6)
                self.add_markdown_text(current_p, line)

    def render_project_block(self, config):
        # Configuration
        title = config.get('title', '')
        items = config.get('items', [])
        tags = config.get('tags', [])
        style = config.get('style', 'left_border') # Default to left border
        border_color_key = config.get('border_color', 'primary_color')
        
        # Colors
        border_color_val = self.resolve_color(border_color_key if border_color_key else 'primary_color')
        
        # Use a Container Table (1x1) for Border/Shading
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Mm(190) # Full width
        
        # Table Properties (Borders, Margins)
        tbl_pr = table._element.tblPr
        
        # Borders
        tblBorders = OxmlElement('w:tblBorders')
        
        # Left Border (Always on for project block usually)
        if style == 'left_border':
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'single')
            left_border.set(qn('w:sz'), '40') # 5px approx
            left_border.set(qn('w:space'), '0')
            left_border.set(qn('w:color'), border_color_val.replace('#', ''))
            tblBorders.append(left_border)
            
        # Add borders to properties
        tbl_pr.append(tblBorders)
        
        # Cell Properties (Padding)
        # Use cell to control content padding
        cell = table.cell(0, 0)
        tcPr = cell._element.get_or_add_tcPr()
        
        # Margins (Padding)
        tcMar = OxmlElement('w:tcMar')
        
        # Top/Bottom Padding (Tight: 2px ~ 30-40 dx equivalent)
        top_mar = OxmlElement('w:top')
        top_mar.set(qn('w:w'), '40') 
        top_mar.set(qn('w:type'), 'dxa')
        
        bottom_mar = OxmlElement('w:bottom')
        bottom_mar.set(qn('w:w'), '40') 
        bottom_mar.set(qn('w:type'), 'dxa')
        
        # Left Padding (15px ~ 225 dxa)
        left_mar = OxmlElement('w:left')
        left_mar.set(qn('w:w'), '225') 
        left_mar.set(qn('w:type'), 'dxa')
        
        tcMar.append(top_mar)
        tcMar.append(bottom_mar)
        tcMar.append(left_mar)
        
        tcPr.append(tcMar)
        
        # Render Content inside Cell
        # 1. Title
        if title:
            p = cell.paragraphs[0]
            self.add_markdown_text(p, title)
            # Find the run and bold/size it (add_markdown_text adds runs)
            # Just enforcing bold on all runs for title
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11) 
            p.paragraph_format.space_after = Pt(3)
        else:
            # If no title, clear the first paragraph or reuse it
            p = cell.paragraphs[0]
            p.clear()

        # 2. Items (Bullets)
        for item in items:
            p = cell.add_paragraph() # New paragraph for item
            p.style = 'List Bullet' # Use default bullet style
            
            # Reduce indentation for bullet
            p_format = p.paragraph_format
            p_format.left_indent = Mm(5)
            p_format.first_line_indent = Mm(-5)
            p_format.space_after = Pt(2)
            
            self.add_markdown_text(p, item)
            
        # 3. Tags (Pills)
        if tags:
            # Add spacer
            p_tags = cell.add_paragraph()
            p_tags.paragraph_format.space_before = Pt(6)
            
            theme_color_hex = self.resolve_color('primary_color').replace('#', '')
            
            for tag in tags:
                # Spacer between pills
                if tag != tags[0]:
                    run_space = p_tags.add_run("  ") 
                
                # Tag Text
                run = p_tags.add_run(f" {tag} ") # Pad with spaces
                run.font.color.rgb = RGBColor.from_string('FFFFFF') # White Text
                run.font.size = self.get_font_size('small', 9)
                run.bold = True
                
                # Shading (Background Color)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), theme_color_hex)
                
                # Append to rPr (Run Properties), NOT the Run itself
                run._element.get_or_add_rPr().append(shd)

        # 4. Add spacer after block
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
    def resolve_color(self, color_name):
        # Same as HTML logic
        if color_name.startswith('#'):
            return color_name
            
        color_map = {
            'primary_color': self.theme.get('primary_color', '#004a99'),
            'accent_color': self.theme.get('accent_color', '#E07000'),
            'text_color': self.theme.get('text_color', '#333333')
        }
        return color_map.get(color_name, '#000000')
    def add_section_title(self, title, style='normal', target=None):
        if target is None:
            target = self.doc
        p = target.add_paragraph()
        
        # If accented, we might want a different look. 
        # For now, let's keep font same but handle underline.
        
        run = p.add_run(title)
        run.bold = True
        run.font.size = self.get_font_size('h1', 16)
        run.font.name = self.theme.get('font_header', 'Arial')
        run.font.color.rgb = hex_to_rgb(get_theme_color(self.theme, 'primary_color'))
       
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        if style == 'accented':
            # Orange underline on text only (mimic span behavior)
            color_hex = get_theme_color(self.theme, 'accent_color').lstrip('#')
            
            # Apply underline to the run
            rPr = run._r.get_or_add_rPr()
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single') # or 'thick'
            u.set(qn('w:color'), color_hex)
            rPr.append(u)
            
        else:
            # Standard Grey Full-Width Border
            # Border Logic
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            pPr.insert(0, pBdr) 
            
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6') # 0.75pt (thinner than accent)
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'EEEEEE')
            
            pBdr.append(bottom)

    def add_markdown_text(self, paragraph, text, font_size_override=None, color_override=None):
        # Improved parser: Handle **Bold** and <color hex="..."></color>
        # We need a regex or state machine. Let's use regex split.
        import re
        
        # Regex to split by bold or color tags
        # Supports: **Bold**  and  <span style="color: #HEX">Text</span>
        # We'll stick to a simpler custom tag for robustness in this regex: [color:#HEX]Text[/color]
        # OR just regex for the standard HTML span we plan to use.
        
        # Pattern: (**.*?**) OR (<span style="color:.*?">.*?</span>)
        pattern = r'(\*\*.*?\*\*|<span style="color:.*?">.*?</span>)'
        parts = re.split(pattern, text)
        
        # Determine base font size
        base_font_size = self.get_font_size('base', 10)
        
        for part in parts:
            if not part: continue
            
            run = paragraph.add_run()
            
            # Defaults
            run.font.name = self.theme.get('font_body', 'Arial')
            run.font.size = font_size_override if font_size_override else base_font_size
            run.font.color.rgb = hex_to_rgb(self.theme.get('text_color', '#333333'))
            
            # 1. Bold
            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
                
            # 2. Color Span
            elif part.startswith('<span') and 'color:' in part:
                # Extract color and text
                # <span style="color: #E65100">Text</span>
                try:
                    color_match = re.search(r'color:\s*(#[0-9a-fA-F]{6})', part)
                    text_match = re.search(r'>(.*?)</span>', part)
                    
                    if color_match and text_match:
                        run.text = text_match.group(1)
                        hex_c = color_match.group(1).lstrip('#')
                        run.font.color.rgb = RGBColor(int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16))
                        # Check if internal text was bold? (Nested not supported by this simple split)
                        if '**' in run.text:
                             # Very hacky: strip ** and make bold. 
                             # Real parser would be recursive.
                             run.text = run.text.replace('**', '')
                             run.bold = True
                    else:
                        run.text = part # Fallback
                except:
                    run.text = part
            
            # 3. Normal
            else:
                run.text = part
            
            # Apply Overrides (Global for this line)
            if font_size_override:
                run.font.size = font_size_override
            if color_override and not part.startswith('<span'): # Don't override explicit spans
                run.font.color.rgb = color_override

    def render_project_block(self, config):
        # 1. Render Items (Highlights)
        style = config.get('style', 'simple')
        items = config.get('items', [])
        tags = config.get('tags', [])
        title = config.get('title') # New: Title inside block
        
        # If left_border style, wrap in table
        if style == 'left_border':
            self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
            table = self.doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Mm(190)
            
            cell = table.cell(0,0)
            target = cell
            
            # Border & Padding
            tcPr = cell._tc.get_or_add_tcPr()
            
            # Padding
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom', 'left', 'right']:
                width = '220' if side == 'left' else '0' # 11pt left padding
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), width)
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
            
            # Border
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
                
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '32')
            left.set(qn('w:space'), '0')
            c_hex = get_theme_color(self.theme, config.get('border_color', 'primary_color'))
            left.set(qn('w:color'), c_hex.lstrip('#'))
            tcBorders.append(left)
            
            # Clear default paragraph
            target.paragraphs[0]._element.getparent().remove(target.paragraphs[0]._element)
            
        else:
            target = self.doc
            
        # 1. Render Title (if exists)
        if title:
            if hasattr(target, 'add_paragraph'):
                p_title = target.add_paragraph()
            else:
                p_title = target.add_paragraph()
            
            p_title.paragraph_format.space_after = Pt(4)
            r = p_title.add_run(title)
            r.bold = True
            r.font.size = Pt(11) # Slightly larger?
            # r.font.color.rgb? Keep black or dark grey as per PDF?
            # PDF image looked black/dark.

        # 2. Render Items (Bullets)
        for item in items:
            if isinstance(item, str):
                details = [item]
            else:
                details = item.get('details', [])
                
            for line in details:
                if hasattr(target, 'add_paragraph'):
                    p = target.add_paragraph()
                else: 
                    p = target.add_paragraph()
                    
                p.style = 'List Bullet'
                if style == 'left_border':
                    p.paragraph_format.space_after = Pt(2)
                    
                self.add_markdown_text(p, line)
                
        # 3. Render Tags (Pills)
        if tags:
            if hasattr(target, 'add_paragraph'):
                p_tags = target.add_paragraph()
            else:
                p_tags = target.add_paragraph()
                
            p_tags.paragraph_format.space_before = Pt(8)
            
            for i, tag in enumerate(tags):
                # Add spacing run between pills
                if i > 0:
                    r_space = p_tags.add_run("  ") # Visual gap
                    r_space.font.size = Pt(9)

                # Pill Run
                # We can't do rounded corners easily in pure DOCX runs without Graphic Objects.
                # But we can do Background Shading + White Text.
                run = p_tags.add_run(f" {tag} ") # Spaces for padding approximation
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255) # White text
                run.bold = True
                
                # Apply Shading (Highlight)
                rPr = run._r.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                # Use primary color for background
                match_color = get_theme_color(self.theme, 'primary_color').lstrip('#')
                shd.set(qn('w:fill'), match_color)
                rPr.append(shd)

        # Spacer after block if it was a table
        if style == 'left_border':
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(12)

    # --- HELPER: HYPERLINK ---
    def add_hyperlink(self, paragraph, url, text, color="#0000FF"):
        # This relationship ID logic remains same
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color.lstrip('#'))
        rPr.append(c)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        
        paragraph._p.append(hyperlink)
