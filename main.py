import yaml
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE
import docx

def load_config(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)
    except FileNotFoundError:
        print(f"Error: Config file not found at {path}")
        sys.exit(1)
    except yaml.YAMLError as exc:
        print(f"Error in configuration file: {exc}")
        sys.exit(1)

def hex_to_rgb(hex_str):
    if not hex_str: return RGBColor(0, 0, 0)
    hex_str = hex_str.lstrip('#')
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def set_page_margins(doc, theme):
    """
    Sets the page margins based on the theme configuration.
    Defaults to 5mm if not specified (very narrow to allow full-width effects).
    """
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(theme.get('Top_margin', 5))
        section.bottom_margin = Mm(theme.get('Bottom_margin', 5))
        section.left_margin = Mm(theme.get('left_margin', 5))
        section.right_margin = Mm(theme.get('right_margin', 5))

def add_stripe_element(doc, config, theme):
    """
    Adds a colored stripe to the top of the BODY of the first page.
    Forces top margin to 0.
    """
    stripe_config = config.get('stripe_element', {})
    if not stripe_config:
        return

    # Check if enabled
    if not stripe_config.get('first_page_only', False):
        return

    section = doc.sections[0]
    
    # 1. Force Top Margin to 0 for the stripe to hit the edge
    # We save the intended top margin to apply as a spacer later if needed
    intended_top_margin = section.top_margin
    section.top_margin = Mm(0)
    
    # 2. Add the stripe logic (Table) to the BODY (doc.add_table)
    
    # Color logic
    color_val = stripe_config.get('color', '000000')
    if color_val == 'primary_color':
        color_hex = theme.get('primary_color', '000000')
    else:
        color_hex = color_val

    # Table wrapper
    page_width = section.page_width
    # Insert table at beginning of document
    # Note: doc.add_table appends. If we want it first, we just call it first in main.
    table = doc.add_table(rows=1, cols=1)
    
    # Full width logic
    left_margin = section.left_margin
    
    table.autofit = False   
    table.allow_autofit = False
    
    # Negative Indent
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._element.insert(0, tblPr)
    
    tblInd = OxmlElement('w:tblInd')
    # section.left_margin is a Length object.
    tblInd.set(qn('w:w'), str(-int(left_margin.twips))) 
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Cell setup
    cell = table.cell(0, 0)
    cell.width = page_width 
    
    # Thickness/Height
    thickness_pt = stripe_config.get('thickness', 5)
    row = table.rows[0]
    row.height = Pt(thickness_pt * 2) 
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Shading
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

    # Zero margins/padding
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), "0")
        node.set(qn('w:type'), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

    # Paragraph spacing zeroing
    p = cell.paragraphs[0]
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)
    p_fmt.line_spacing = 1.0 
    run = p.runs[0] if p.runs else p.add_run()
    run.font.size = Pt(1)

    # 3. Add Spacer Paragraph to restore "visual" margin
    # The stripe takes up some space (thickness). 
    # But usually we want the text to start at 'Top_margin'.
    # So we add a paragraph with space_after = intended_top_margin (approx)
    # spacer = doc.add_paragraph()
    # spacer.paragraph_format.space_before = Pt(0)
    # spacer.paragraph_format.space_after = Pt(0)
    # spacer.paragraph_format.line_spacing = Pt(0) # Minimal height
    # # We use space_before on the Next element or just a fixed spacer height?
    # # Let's convert intended margin (EMU) to Points
    # margin_pt = intended_top_margin.pt
    # spacer.paragraph_format.space_after = Pt(margin_pt)
    
    # Better: Just let the user's TEXT_ELEMENT handle its own placement?
    # User said "handle the spacing".
    # I will add a spacer paragraph equal to the Intended Top Margin.
    # But we subtracted 0, so effectively we are at 0 + StripeHeight.
    # Text should start at IntendedTopMargin.
    # Use exact spacing.
    
    spacer = doc.add_paragraph()
    run_s = spacer.add_run()
    run_s.font.size = Pt(1) # minimize text height
    # Set exact spacing
    # spacer.paragraph_format.space_after = Pt(intended_top_margin.pt) 
    # actually space_before is better for the text, but a separate spacer is cleaner.
    spacer.paragraph_format.line_spacing = Pt(intended_top_margin.pt) # Force height?
    # safest is space_after
    spacer.paragraph_format.space_after = Pt(intended_top_margin.pt)
    spacer.paragraph_format.space_before = Pt(0)
    # Remove line spacing influence
    spacer.paragraph_format.line_spacing_rule = WD_ROW_HEIGHT_RULE.EXACTLY 
    # wait paragraph format uses WD_LINE_SPACING.EXACTLY if passed as Pt? No. 
    # docx ParagraphFormat.line_spacing set to a Length (Pt) implies EXACTLY.
    spacer.paragraph_format.line_spacing = Pt(0) 
    # Wait, if line spacing is 0, space_after does the job.


def add_text_element(doc, config, theme):
    """
    Adds a text element based on configuration.
    """
    text_config = config.get('TEXT_ELEMENT', {})
    if not text_config:
        return

    text = text_config.get('words', '')
    if not text:
        return

    paragraph = doc.add_paragraph()
    
    # alignment
    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    # Handle user typo in yaml 'font_alightnment' or correct 'font_alignment'
    alignment_str = text_config.get('font_alightnment', text_config.get('font_alignment', 'left')).lower()
    paragraph.alignment = align_map.get(alignment_str, WD_ALIGN_PARAGRAPH.LEFT)

    run = paragraph.add_run(text)

    # Size
    size = text_config.get('font_size')
    if size:
        run.font.size = Pt(size)

    # Style
    style = text_config.get('font_style', '').lower()
    if 'bold' in style:
        run.bold = True
    if 'italic' in style:
        run.italic = True

    # Color
    color_val = text_config.get('font_color', '000000')
    if color_val == 'primary_color':
        color_hex = theme.get('primary_color', '000000')
    else:
        color_hex = color_val.lstrip('#') # Ensure clean hex
        
    run.font.color.rgb = hex_to_rgb(color_hex)

def add_hyperlink(paragraph, url, text, color_rgb, underline=False):
    """
    A helper to add a hyperlink to a paragraph.
    """
    # This gets access to the document.part.relate_to
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color if provided
    if color_rgb:
        # color_rgb is an RGBColor object, we need hex string without #
        # RGBColor doesn't easily give hex, but we can format it
        # Actually docx RGBColor can be converted to str? No.
        # We can construct the hex string if we have the components
        # But wait, color_rgb passed here might be the RGBColor object from main.
        # Let's extract hex.
        # simpler: just pass the hex string to this function
        pass
    
    # We will handle styling manually using the run properties we append
    
    # Custom styling implementation for the run inside hyperlink
    # Join the formatting
    if color_rgb:
        c = OxmlElement('w:color')
        # Assuming color_rgb is a string "RRGGBB" for this low level xml
        c.set(qn('w:val'), color_rgb) 
        rPr.append(c)
    
    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)
    else:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def add_compound_text_element(doc, config, theme):
    """
    Adds a text element composed of multiple items with optional links.
    """
    comp_config = config.get('COMPOUND_TEXT_ELEMENT', {})
    if not comp_config:
        return

    items = comp_config.get('items', [])
    if not items:
        return

    paragraph = doc.add_paragraph()
    
    # alignment
    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    alignment_str = comp_config.get('font_alignment', 'center').lower()
    paragraph.alignment = align_map.get(alignment_str, WD_ALIGN_PARAGRAPH.CENTER)

    # Common style props
    font_size = comp_config.get('font_size', 11)
    color_val = comp_config.get('font_color', '000000')
    if color_val == 'primary_color':
        color_hex = theme.get('primary_color', '000000')
    else:
        color_hex = color_val.lstrip('#')
    
    separator = comp_config.get('separator', ' • ')
    
    for i, item in enumerate(items):
        text = item.get('text', '')
        link = item.get('link')
        
        # Add separator if not first
        if i > 0:
            run_sep = paragraph.add_run(separator)
            run_sep.font.size = Pt(font_size)
            run_sep.font.color.rgb = hex_to_rgb(color_hex)
        
        if link:
            # Item specific color?
            item_color_val = item.get('font_color')
            if item_color_val:
                if item_color_val == 'primary_color':
                    item_hex = theme.get('primary_color', '000000')
                else:
                    item_hex = item_color_val.lstrip('#')
                # Use item specific color
                add_hyperlink_styled(paragraph, link, text, item_hex, font_size)
            else:
                # Use default element color
                add_hyperlink_styled(paragraph, link, text, color_hex, font_size)
        else:
            run = paragraph.add_run(text)
            run.font.size = Pt(font_size)
            run.font.color.rgb = hex_to_rgb(color_hex)

    # Post-process for size on hyperlinks? 
    # It's hard to target them easily.
    # Better to control size in the xml creation.

def add_hyperlink_styled(paragraph, url, text, color_hex, size_pt, font_name=None):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Color
    if color_hex:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color_hex)
        rPr.append(c)
    
    # Size
    if size_pt:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2))) # half-points
        rPr.append(sz)
        
    # Font
    if font_name:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def main(config_file, output_file):
    print(f"Loading configuration from {config_file}...")
    config = load_config(config_file)
    theme = config.get('theme', {})
    
    doc = Document()
    
    set_page_margins(doc, theme)
    
    # Process elements
    add_stripe_element(doc, config, theme)
    add_text_element(doc, config, theme)
    add_compound_text_element(doc, config, theme)
    
    # Add some dummy content to body to see the page
    doc.add_paragraph("Resume Content Starts Here...")
    
    doc.save(output_file)
    print(f"Document generated: {output_file}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        # Default for dev
        config_path = "test.yaml"
    else:
        config_path = sys.argv[1]
        
    main(config_path, "output.docx")
